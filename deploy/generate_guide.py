from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper functions
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x25, 0x1E)
        run.font.size = Pt(26)

def add_h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x5C, 0x3D)
        run.font.size = Pt(18)

def add_h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2A, 0x3A, 0x2C)
        run.font.size = Pt(14)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run("ВАЖНО: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return p

def add_tip(text):
    p = doc.add_paragraph()
    run = p.add_run("СОВЕТ: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x77, 0x00)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x00, 0x77, 0x00)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x88)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ==========================================
# DOCUMENT START
# ==========================================

add_title("QR Restoran")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Полная инструкция по установке и запуску")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("от регистрации до рабочего приложения")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# ==========================================
# TABLE OF CONTENTS
# ==========================================
add_h1("Содержание")
add_para("Часть 1: Что нам понадобится")
add_para("Часть 2: Регистрация в Oracle Cloud (бесплатно)")
add_para("Часть 3: Создание виртуального сервера")
add_para("Часть 4: Настройка доступа (открытие портов)")
add_para("Часть 5: Подключение к серверу через SSH")
add_para("Часть 6: Установка Docker на сервер")
add_para("Часть 7: Загрузка проекта на сервер")
add_para("Часть 8: Настройка переменных окружения")
add_para("Часть 9: Запуск приложения")
add_para("Часть 10: Проверка работы")
add_para("Часть 11: Привязка домена (необязательно)")
add_para("Часть 12: SSL сертификат (HTTPS)")
add_para("Часть 13: Обновление приложения")
add_para("Часть 14: Решение проблем")
add_para("Приложение: Полезные команды")
doc.add_page_break()

# ==========================================
# PART 1
# ==========================================
add_h1("Часть 1: Что нам понадобится")

add_para("Прежде чем начать, убедитесь, что у вас есть:")
add_numbered("Компьютер с Windows, Mac или Linux")
add_numbered("Стабильное интернет-соединение")
add_numbered("Банковская карта (для регистрации в Oracle Cloud — деньги НЕ спишут)")
add_numbered("Электронная почта (email)")
add_numbered("Около 1-2 часов свободного времени")

doc.add_paragraph()
add_para("Что такое Oracle Cloud Free Tier?", bold=True)
add_para("Oracle Cloud — это облачная платформа (как Amazon AWS или Google Cloud). У них есть бесплатный тариф (Free Tier), который даёт вам виртуальный сервер НАВСЕГДА бесплатно. На этом сервере мы и запустим наше приложение.")

add_para("Что вы получите бесплатно:", bold=True)
add_bullet("2 виртуальных сервера (1 ГБ оперативной памяти каждый)")
add_bullet("200 ГБ хранилища")
add_bullet("10 ТБ/месяц трафика")

add_warning("Банковская карта нужна только для верификации. Oracle не будет списывать деньги, пока вы используете бесплатные ресурсы. Это стандартная практика всех облачных провайдеров.")

doc.add_page_break()

# ==========================================
# PART 2
# ==========================================
add_h1("Часть 2: Регистрация в Oracle Cloud")

add_h2("Шаг 2.1: Открытие сайта")
add_numbered("Откройте браузер (Chrome, Firefox или любой другой)")
add_numbered("Перейдите по адресу: https://cloud.oracle.com/free")
add_numbered('Нажмите кнопку "Start for free" (Начать бесплатно)')

add_h2("Шаг 2.2: Заполнение формы регистрации")
add_numbered('В поле "Country" выберите вашу страну')
add_numbered('В поле "Name" введите ваше имя латинскими буквами')
add_numbered('В поле "Email" введите вашу электронную почту')
add_numbered('Придумайте пароль (минимум 12 символов, должен содержать заглавную букву, строчную букву и цифру)')
add_numbered('Нажмите "Verify my email"')

add_h2("Шаг 2.3: Подтверждение email")
add_numbered("Зайдите в вашу электронную почту")
add_numbered("Найдите письмо от Oracle (может быть в папке Спам)")
add_numbered("Нажмите на ссылку подтверждения в письме")

add_h2("Шаг 2.4: Выбор региона")
add_para("После подтверждения email вам предложат выбрать регион. Это физическое расположение вашего сервера.")
add_warning('Регион нельзя изменить после регистрации! Выбирайте ближайший к вам.')
add_para("Рекомендуемые регионы:")
add_bullet("Для Азербайджана/Турции: Germany Central (Frankfurt) или UAE East (Dubai)")
add_bullet("Для СНГ: Germany Central (Frankfurt)")

add_h2("Шаг 2.5: Привязка банковской карты")
add_numbered("Введите данные вашей банковской карты")
add_numbered("Oracle спишет и вернёт $1 для проверки карты")
add_numbered('После проверки нажмите "Start my free trial"')

add_tip("Если карту не принимают, попробуйте другую карту (Visa/Mastercard). Карты Мир и некоторые дебетовые карты могут не приниматься.")

add_para("Подождите 5-10 минут, пока Oracle создаст ваш аккаунт. Вы получите email, когда всё будет готово.")

doc.add_page_break()

# ==========================================
# PART 3
# ==========================================
add_h1("Часть 3: Создание виртуального сервера")

add_h2("Шаг 3.1: Вход в Oracle Cloud Console")
add_numbered("Перейдите на https://cloud.oracle.com")
add_numbered("Введите имя вашего аккаунта (Cloud Account Name)")
add_numbered("Нажмите Next, затем введите email и пароль")

add_h2("Шаг 3.2: Создание SSH ключа")
add_para("SSH ключ — это как электронный пароль для безопасного подключения к серверу. Его нужно создать на вашем компьютере.", bold=True)

add_para("Для Windows:", bold=True)
add_numbered('Нажмите клавиши Win+R, введите "cmd" и нажмите Enter')
add_numbered("В открывшемся окне введите команду:")
add_code("ssh-keygen -t rsa -b 4096")
add_numbered('Нажмите Enter 3 раза (не вводите пароль, просто жмите Enter)')
add_numbered("Ключ сохранится в папку C:\\Users\\ВАШЕ_ИМЯ\\.ssh\\")
add_numbered("Запомните расположение файлов:")
add_bullet("id_rsa — это ваш ПРИВАТНЫЙ ключ (никому не показывайте!)", bold_prefix="")
add_bullet("id_rsa.pub — это ПУБЛИЧНЫЙ ключ (его мы загрузим в Oracle)", bold_prefix="")

add_para("Для Mac/Linux:", bold=True)
add_numbered("Откройте Терминал (Terminal)")
add_numbered("Введите команду:")
add_code("ssh-keygen -t rsa -b 4096")
add_numbered("Нажмите Enter 3 раза")
add_numbered("Ключи сохранятся в ~/.ssh/")

add_h2("Шаг 3.3: Создание виртуальной машины (VM)")
add_numbered('В Oracle Cloud Console нажмите на меню (три полоски слева вверху)')
add_numbered('Выберите "Compute" → "Instances"')
add_numbered('Нажмите синюю кнопку "Create Instance"')

add_para("Заполните параметры:", bold=True)
add_numbered('Name: введите "qr-restoran"')
add_numbered("Placement: оставьте как есть")

add_numbered("Image and shape — ОЧЕНЬ ВАЖНО:", bold_prefix="")
add_bullet('Нажмите "Edit" рядом с Image and shape')
add_bullet('В разделе Image: выберите "Canonical Ubuntu" версии 22.04')
add_bullet('В разделе Shape: нажмите "Change shape"')
add_bullet('Выберите "VM.Standard.E2.1.Micro" — это бесплатный вариант!')
add_bullet('Убедитесь, что написано "Always Free-eligible"')

add_numbered("Networking: оставьте как есть (Oracle создаст сеть автоматически)")

add_numbered("Add SSH keys:", bold_prefix="")
add_bullet('Выберите "Upload public key files (.pub)"')
add_bullet('Нажмите "Browse" и выберите файл id_rsa.pub')
add_bullet("Windows: C:\\Users\\ВАШЕ_ИМЯ\\.ssh\\id_rsa.pub")
add_bullet("Mac/Linux: ~/.ssh/id_rsa.pub")

add_numbered("Boot volume: оставьте как есть (50 ГБ)")

add_numbered('Нажмите "Create"!')

add_para("Подождите 2-5 минут, пока сервер создаётся. Когда значок станет зелёным (RUNNING), сервер готов.")

add_h2("Шаг 3.4: Запомните IP-адрес!")
add_para('На странице вашего сервера найдите "Public IP Address". Это будет что-то вроде: 129.213.XX.XX')
add_warning("Запишите этот IP-адрес! Он понадобится на каждом следующем шаге.")

doc.add_page_break()

# ==========================================
# PART 4
# ==========================================
add_h1("Часть 4: Настройка доступа (открытие портов)")

add_para("По умолчанию Oracle блокирует весь входящий трафик. Нам нужно открыть порты 80 (HTTP) и 443 (HTTPS), чтобы люди могли заходить на сайт.")

add_h2("Шаг 4.1: Открытие портов в Oracle Security List")
add_numbered('В меню Oracle Cloud: "Networking" → "Virtual Cloud Networks"')
add_numbered("Нажмите на вашу VCN (она была создана автоматически)")
add_numbered('Слева нажмите "Security Lists"')
add_numbered('Нажмите на "Default Security List..."')
add_numbered('Нажмите "Add Ingress Rules"')

add_para("Добавьте правило для HTTP (порт 80):", bold=True)
add_bullet("Source Type: CIDR")
add_bullet("Source CIDR: 0.0.0.0/0")
add_bullet("IP Protocol: TCP")
add_bullet("Destination Port Range: 80")
add_bullet('Нажмите "Add Ingress Rules"')

add_para("Повторите для HTTPS (порт 443):", bold=True)
add_bullet('Снова нажмите "Add Ingress Rules"')
add_bullet("Source CIDR: 0.0.0.0/0")
add_bullet("IP Protocol: TCP")
add_bullet("Destination Port Range: 443")
add_bullet('Нажмите "Add Ingress Rules"')

doc.add_page_break()

# ==========================================
# PART 5
# ==========================================
add_h1("Часть 5: Подключение к серверу через SSH")

add_para("Теперь нам нужно подключиться к нашему серверу, чтобы установить на нём программы.")

add_h2("Для Windows:")
add_numbered('Нажмите Win+R, введите "cmd" и нажмите Enter')
add_numbered("Введите команду (замените IP на ваш IP-адрес):")
add_code("ssh ubuntu@ВАШ_IP_АДРЕС")
add_para("Пример:")
add_code("ssh ubuntu@129.213.45.67")
add_numbered('Если спросит "Are you sure you want to continue?" — напишите "yes" и нажмите Enter')
add_numbered("Вы подключены к серверу! Вы увидите приглашение вида: ubuntu@qr-restoran:~$")

add_h2("Для Mac/Linux:")
add_numbered("Откройте Терминал")
add_numbered("Введите команду:")
add_code("ssh ubuntu@ВАШ_IP_АДРЕС")
add_numbered('Напишите "yes" если спросит')

add_warning("Если подключение не работает:")
add_bullet("Проверьте, что IP-адрес правильный")
add_bullet("Проверьте, что сервер имеет статус RUNNING в Oracle Console")
add_bullet("Подождите 5 минут после создания сервера")
add_bullet("Убедитесь, что файл id_rsa находится в папке .ssh/")

doc.add_page_break()

# ==========================================
# PART 6
# ==========================================
add_h1("Часть 6: Установка Docker на сервер")

add_para("Docker — это программа, которая позволяет запускать приложения в изолированных контейнерах. Это как виртуальный компьютер внутри вашего сервера.")

add_h2("Шаг 6.1: Обновление системы")
add_para("Введите эти команды по одной (копируйте и вставляйте):")
add_code("sudo apt-get update")
add_para("Подождите завершения (может занять 1-2 минуты).")
add_code("sudo apt-get upgrade -y")
add_para("Подождите завершения (может занять 3-5 минут).")

add_h2("Шаг 6.2: Установка Docker")
add_code("curl -fsSL https://get.docker.com | sh")
add_para("Подождите 2-3 минуты.")

add_h2("Шаг 6.3: Настройка прав доступа")
add_code("sudo usermod -aG docker $USER")

add_h2("Шаг 6.4: Установка Docker Compose")
add_code("sudo apt-get install -y docker-compose-plugin")

add_h2("Шаг 6.5: Перезайдите на сервер")
add_para("Это нужно, чтобы новые права вступили в силу:")
add_code("exit")
add_code("ssh ubuntu@ВАШ_IP_АДРЕС")

add_h2("Шаг 6.6: Проверка установки")
add_code("docker --version")
add_para("Должно показать что-то вроде: Docker version 24.x.x")
add_code("docker compose version")
add_para("Должно показать что-то вроде: Docker Compose version v2.x.x")

add_h2("Шаг 6.7: Открытие портов в файрволе Ubuntu")
add_para("Oracle Ubuntu имеет встроенный файрвол. Его тоже нужно настроить:")
add_code("sudo iptables -I INPUT -p tcp --dport 80 -j ACCEPT")
add_code("sudo iptables -I INPUT -p tcp --dport 443 -j ACCEPT")
add_code("sudo apt-get install -y iptables-persistent")
add_para("Когда спросит 'Save current rules?' — выберите Yes (оба раза).")
add_code("sudo netfilter-persistent save")

add_tip("Если iptables-persistent уже установлен и спрашивает о сохранении — нажимайте Yes.")

doc.add_page_break()

# ==========================================
# PART 7
# ==========================================
add_h1("Часть 7: Загрузка проекта на сервер")

add_h2("Способ A: Через GitHub (рекомендуется)")

add_para("Если проект сохранён на GitHub:", bold=True)
add_numbered("Установите git (если ещё не установлен):")
add_code("sudo apt-get install -y git")
add_numbered("Клонируйте репозиторий:")
add_code("git clone https://github.com/ВАШЕ_ИМЯ/ВАШ_РЕПОЗИТОРИЙ.git")
add_numbered("Перейдите в папку проекта:")
add_code("cd ВАШ_РЕПОЗИТОРИЙ")

add_h2("Способ B: Загрузка файлов напрямую (если нет GitHub)")

add_para("Если у вас файлы проекта на компьютере:", bold=True)
add_numbered("На вашем ЛОКАЛЬНОМ компьютере (не на сервере!) откройте новое окно командной строки")
add_numbered("Перейдите в папку с проектом:")
add_code("cd путь/к/вашему/проекту")
add_numbered("Скопируйте файлы на сервер:")
add_code("scp -r . ubuntu@ВАШ_IP:~/qr-restoran/")
add_numbered("На сервере перейдите в папку:")
add_code("cd ~/qr-restoran")

doc.add_page_break()

# ==========================================
# PART 8
# ==========================================
add_h1("Часть 8: Настройка переменных окружения")

add_para("Переменные окружения — это настройки, которые приложение читает при запуске (адреса, пароли и т.д.)")

add_h2("Шаг 8.1: Перейдите в папку deploy")
add_code("cd deploy")

add_h2("Шаг 8.2: Создайте файл .env")
add_code("cp .env.example .env")

add_h2("Шаг 8.3: Отредактируйте файл")
add_code("nano .env")

add_para("Измените содержимое на:", bold=True)
add_code("REACT_APP_BACKEND_URL=http://ВАШ_IP_АДРЕС")
add_code("MONGO_URL=mongodb://mongo:27017")
add_code("DB_NAME=restaurant_db")

add_para("Пример (если ваш IP — 129.213.45.67):", bold=True)
add_code("REACT_APP_BACKEND_URL=http://129.213.45.67")
add_code("MONGO_URL=mongodb://mongo:27017")
add_code("DB_NAME=restaurant_db")

add_para("Как сохранить файл в nano:", bold=True)
add_numbered("Нажмите Ctrl+O (сохранить)")
add_numbered("Нажмите Enter (подтвердить имя файла)")
add_numbered("Нажмите Ctrl+X (выйти из редактора)")

add_warning("Замените ВАШ_IP_АДРЕС на реальный IP вашего сервера из Части 3!")

doc.add_page_break()

# ==========================================
# PART 9
# ==========================================
add_h1("Часть 9: Запуск приложения")

add_h2("Шаг 9.1: Запуск")
add_para("Убедитесь, что вы в папке deploy:")
add_code("cd ~/qr-restoran/deploy")
add_para("ИЛИ (если клонировали через git)")
add_code("cd ~/ВАШ_РЕПОЗИТОРИЙ/deploy")

add_para("Запустите приложение:", bold=True)
add_code("docker compose up -d --build")

add_para("Это займёт 5-10 минут при первом запуске. Docker будет:")
add_bullet("Скачивать необходимые образы (Python, Node.js, MongoDB, Nginx)")
add_bullet("Устанавливать зависимости")
add_bullet("Собирать React-приложение")
add_bullet("Запускать все сервисы")

add_h2("Шаг 9.2: Проверка статуса")
add_para("Когда команда завершится, проверьте что всё работает:")
add_code("docker compose ps")

add_para("Вы должны увидеть два контейнера со статусом 'Up':")
add_bullet("deploy-app-1 — само приложение (backend + frontend)")
add_bullet("deploy-mongo-1 — база данных MongoDB")

add_h2("Шаг 9.3: Просмотр логов (если что-то не работает)")
add_code("docker compose logs -f")
add_para("Нажмите Ctrl+C чтобы выйти из просмотра логов.")

doc.add_page_break()

# ==========================================
# PART 10
# ==========================================
add_h1("Часть 10: Проверка работы")

add_h2("Шаг 10.1: Откройте браузер")
add_para("На вашем компьютере откройте браузер и перейдите по адресу:")
add_code("http://ВАШ_IP_АДРЕС")
add_para("Пример: http://129.213.45.67")

add_h2("Шаг 10.2: Первый вход")
add_para("Вы увидите страницу входа. Введите:", bold=True)
add_bullet("Имя пользователя: ", bold_prefix="owner")
add_bullet("Пароль: ", bold_prefix="owner123")

add_h2("Шаг 10.3: Проверьте основные функции")
add_numbered("Зайдите в панель администратора")
add_numbered("Создайте тестовое заведение (venue)")
add_numbered("Добавьте столы")
add_numbered("Добавьте категории меню и блюда")
add_numbered("Попробуйте отсканировать QR-код или открыть ссылку для клиента")

add_para("Поздравляем! Ваше приложение работает!", bold=True)

doc.add_page_break()

# ==========================================
# PART 11
# ==========================================
add_h1("Часть 11: Привязка домена (необязательно)")

add_para("Если вы хотите, чтобы ваш сайт был доступен по красивому адресу (например, restoran.az вместо 129.213.45.67), вам нужен домен.")

add_h2("Шаг 11.1: Покупка домена")
add_para("Купите домен у любого регистратора:")
add_bullet("Namecheap.com — от $8/год")
add_bullet("GoDaddy.com — от $10/год")
add_bullet("Azeri регистраторы для .az доменов")

add_h2("Шаг 11.2: Настройка DNS")
add_numbered("Зайдите в панель управления вашего домена")
add_numbered('Найдите раздел "DNS Records" или "DNS Management"')
add_numbered("Создайте запись:")
add_bullet("Тип: A")
add_bullet("Имя: @ (или пустое)")
add_bullet("Значение: ВАШ_IP_АДРЕС (IP вашего Oracle сервера)")
add_bullet("TTL: Automatic или 300")
add_numbered("Если хотите www:")
add_bullet("Тип: A")
add_bullet("Имя: www")
add_bullet("Значение: ВАШ_IP_АДРЕС")
add_numbered("Подождите 15-30 минут пока DNS обновится")

add_h2("Шаг 11.3: Обновите REACT_APP_BACKEND_URL")
add_para("После привязки домена обновите настройки:")
add_code("cd ~/qr-restoran/deploy")
add_code("nano .env")
add_para("Измените строку:")
add_code("REACT_APP_BACKEND_URL=http://вашдомен.com")
add_para("Перезапустите приложение:")
add_code("docker compose up -d --build")

doc.add_page_break()

# ==========================================
# PART 12
# ==========================================
add_h1("Часть 12: SSL сертификат (HTTPS)")

add_para("SSL сертификат нужен, чтобы ваш сайт работал по HTTPS (с замочком в браузере). Это бесплатно с помощью Let's Encrypt.")

add_warning("Для SSL обязательно нужен домен! С IP-адресом SSL не работает.")

add_h2("Шаг 12.1: Установка Certbot")
add_para("Подключитесь к серверу и выполните:")
add_code("sudo apt-get install -y certbot")

add_h2("Шаг 12.2: Остановите приложение (временно)")
add_code("cd ~/qr-restoran/deploy")
add_code("docker compose down")

add_h2("Шаг 12.3: Получение сертификата")
add_code("sudo certbot certonly --standalone -d вашдомен.com")
add_para("Введите email и согласитесь с условиями.")

add_h2("Шаг 12.4: Обновите настройки")
add_para("Измените .env:")
add_code("nano .env")
add_code("REACT_APP_BACKEND_URL=https://вашдомен.com")

add_h2("Шаг 12.5: Перезапустите приложение")
add_code("docker compose up -d --build")

add_tip("Сертификаты Let's Encrypt действуют 90 дней. Для автоматического обновления добавьте в crontab:")
add_code('sudo crontab -e')
add_para("Добавьте строку:")
add_code("0 0 1 * * certbot renew --quiet")

doc.add_page_break()

# ==========================================
# PART 13
# ==========================================
add_h1("Часть 13: Обновление приложения")

add_para("Когда вы получите новую версию приложения, вот как её установить:")

add_h2("Если используете GitHub:")
add_code("cd ~/qr-restoran")
add_code("git pull")
add_code("cd deploy")
add_code("docker compose up -d --build")

add_h2("Если загружаете файлы вручную:")
add_numbered("На ЛОКАЛЬНОМ компьютере:")
add_code("scp -r . ubuntu@ВАШ_IP:~/qr-restoran/")
add_numbered("На сервере:")
add_code("cd ~/qr-restoran/deploy")
add_code("docker compose up -d --build")

doc.add_page_break()

# ==========================================
# PART 14
# ==========================================
add_h1("Часть 14: Решение проблем")

add_h2("Проблема: Сайт не открывается")
add_numbered("Проверьте, что контейнеры запущены:")
add_code("docker compose ps")
add_numbered("Проверьте логи:")
add_code("docker compose logs app")
add_numbered("Проверьте файрвол Oracle (Часть 4)")
add_numbered("Проверьте файрвол Ubuntu:")
add_code("sudo iptables -L -n | grep 80")
add_para("Если нет правила для порта 80:")
add_code("sudo iptables -I INPUT -p tcp --dport 80 -j ACCEPT")
add_code("sudo netfilter-persistent save")

add_h2("Проблема: Не хватает памяти (RAM)")
add_para("Oracle Free Tier даёт всего 1 ГБ RAM. Если приложение работает медленно или падает, добавьте swap-файл:")
add_code("sudo fallocate -l 2G /swapfile")
add_code("sudo chmod 600 /swapfile")
add_code("sudo mkswap /swapfile")
add_code("sudo swapon /swapfile")
add_code("echo '/swapfile none swap sw 0 0' | sudo tee -a /etc/fstab")

add_h2("Проблема: MongoDB не запускается")
add_code("docker compose logs mongo")
add_para("Если ошибка связана с местом на диске:")
add_code("docker system prune -a")
add_para("Это удалит неиспользуемые образы и освободит место.")

add_h2("Проблема: Не могу подключиться по SSH")
add_bullet("Подождите 5 минут после создания сервера")
add_bullet("Проверьте, что используете правильный IP")
add_bullet("Убедитесь, что файл ключа id_rsa в папке .ssh/")
add_bullet("Для Windows попробуйте PuTTY (бесплатная программа для SSH)")

add_h2("Проблема: docker compose не найден")
add_para("Попробуйте с дефисом:")
add_code("docker-compose up -d --build")
add_para("Если не работает, переустановите:")
add_code("sudo apt-get install -y docker-compose-plugin")

doc.add_page_break()

# ==========================================
# APPENDIX
# ==========================================
add_h1("Приложение: Полезные команды")

add_h2("Управление приложением")
p = doc.add_paragraph()
add_bullet("Запуск: ")
add_code("docker compose up -d --build")
add_bullet("Остановка: ")
add_code("docker compose down")
add_bullet("Перезапуск: ")
add_code("docker compose restart")
add_bullet("Статус: ")
add_code("docker compose ps")
add_bullet("Логи (все): ")
add_code("docker compose logs -f")
add_bullet("Логи (только backend): ")
add_code("docker compose logs -f app")
add_bullet("Логи (только MongoDB): ")
add_code("docker compose logs -f mongo")

add_h2("Резервное копирование базы данных")
add_para("Создание бэкапа:", bold=True)
add_code("docker compose exec mongo mongodump --out /data/backup")
add_code("docker cp $(docker compose ps -q mongo):/data/backup ./backup_$(date +%Y%m%d)")

add_para("Восстановление из бэкапа:", bold=True)
add_code("docker cp ./backup_ДАТА $(docker compose ps -q mongo):/data/backup")
add_code("docker compose exec mongo mongorestore /data/backup")

add_h2("Мониторинг сервера")
add_bullet("Использование диска: ")
add_code("df -h")
add_bullet("Использование памяти: ")
add_code("free -h")
add_bullet("Запущенные процессы: ")
add_code("htop")

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("QR Restoran — Система управления рестораном")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = "/app/deploy/QR_Restoran_Installation_Guide_RU.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
